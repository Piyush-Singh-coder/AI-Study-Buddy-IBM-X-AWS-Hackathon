from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

def create_sample_paper_docx(paper_data: dict) -> BytesIO:
    """Generates a DOCX file from the sample paper data."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Generated Sample Paper', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("Based on uploaded Previous Year Questions pattern")
    doc.add_paragraph('_' * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacer

    paper = paper_data.get('paper', [])
    if not paper:
        doc.add_paragraph("No questions generated.")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    for section in paper:
        heading = doc.add_heading(f"{section['section']} ({section.get('marks', '?')} Marks each)", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        questions = section.get('questions', [])
        for i, q in enumerate(questions, 1):
            # Question Paragraph
            p = doc.add_paragraph()
            run = p.add_run(f"Q{i}. {q.get('question', '')}")
            run.bold = True
            
            # Answer Paragraph (indented slightly or italicized)
            if 'answer' in q:
                ans_p = doc.add_paragraph(f"Answer: {q['answer']}")
                ans_p.style = 'No Spacing'
                ans_p.paragraph_format.left_indent = Inches(0.5)
                ans_p.runs[0].font.italic = True
                ans_p.runs[0].font.color.rgb = None # Default
            
            doc.add_paragraph() # Spacer between questions

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
